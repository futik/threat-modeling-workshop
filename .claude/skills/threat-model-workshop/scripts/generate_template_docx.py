#!/usr/bin/env python3
"""
Generate the participant workshop document as a single Google-Docs-friendly .docx.

The document combines what a team needs during the exercise in one file, as two
sections separated by a page break so they read like tabs:
  1. Product & architecture — the NeuroScan 3000 scenario (from scenario/*.md)
  2. Threat model worksheet — the fillable Q1–Q4 worksheet

The workshop *introduction* is intentionally NOT included here — it lives in
workshop/00-introduction.md and is shared from the repo by the facilitator.

The architecture diagram is rendered as a PNG image (via matplotlib) and embedded,
rather than drawn with ASCII box characters. ASCII art relies on a monospaced
font and exact column alignment, which breaks on Google Docs import; an embedded
image renders identically in Word and Google Docs.

Design constraints for clean Google Docs import:
- Built-in heading styles only (Heading 1/2/3) so Google Docs keeps the outline.
- Standard font (Calibri) and simple tables with a bold header row.
- No content controls, form fields, macros, or text boxes (Google Docs strips them).
- Blank table cells / underscore lines act as fill-in areas.
- Diagrams are embedded images (survive import intact).
- Page break between the two sections. Google Docs has no way to embed real
  document "tabs" via .docx import; the page-break sections plus the Heading 1
  outline are the portable equivalent. Teams who want literal Google Docs tabs
  can split the two Heading 1 sections into tabs after uploading (see SKILL.md).

Requirements: python-docx and matplotlib.

Usage:
    python generate_template_docx.py [output_path]
Default output: ../../../templates/threat-model-template.docx (relative to this script)
"""

import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x49, 0x7D)      # dark blue for headings
HEADER_FILL = "1F497D"                     # table header shading
NOTE = RGBColor(0x55, 0x55, 0x55)          # grey for hint/italic notes


def set_cell_background(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for lvl, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        st = doc.styles[lvl]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = ACCENT


def add_hint(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = NOTE
    r.font.size = Pt(10)
    return p


def add_table(doc, headers, n_blank_rows, example_row=None):
    rows = 1 + (1 if example_row else 0) + n_blank_rows
    table = doc.add_table(rows=rows, cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_background(cell, HEADER_FILL)
    r_idx = 1
    if example_row:
        for j, val in enumerate(example_row):
            cell = table.rows[r_idx].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.italic = True
            run.font.color.rgb = NOTE
            run.font.size = Pt(10)
        r_idx += 1
    doc.add_paragraph()
    return table


def blank_lines(doc, n=3):
    for _ in range(n):
        doc.add_paragraph("_______________________________________________________________")


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_info_table(doc, headers, rows):
    """A plain content table (not a fill-in worksheet table): header row + data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_background(cell, HEADER_FILL)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def render_architecture_png():
    """Render the NeuroScan 3000 architecture as a PNG and return its path.

    Uses matplotlib so there is no system-level graphviz dependency. Returns a
    path in a temp dir; the caller embeds it into the .docx (python-docx copies
    the bytes into the package, so the temp file can be discarded afterwards).
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyArrow, FancyBboxPatch

    accent = "#1F497D"
    box_fill = "#EAF0F7"
    zone_edge = "#1F497D"

    fig, ax = plt.subplots(figsize=(9, 7))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")

    def zone(x, y, w, h, label):
        ax.add_patch(FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.4,rounding_size=1.5",
            linewidth=1.5, edgecolor=zone_edge, facecolor="none", linestyle="--"))
        ax.text(x + 1.5, y + h - 2.5, label, fontsize=10, fontweight="bold",
                color=accent, va="top")

    def box(cx, cy, w, h, lines, external=False):
        # external=True renders a dashed, greyed box for out-of-scope systems
        edge = "#888888" if external else accent
        fill = "#F2F2F2" if external else box_fill
        ax.add_patch(FancyBboxPatch(
            (cx - w / 2, cy - h / 2), w, h,
            boxstyle="round,pad=0.2,rounding_size=1.0",
            linewidth=1.2, edgecolor=edge, facecolor=fill,
            linestyle="--" if external else "-"))
        ax.text(cx, cy, "\n".join(lines), fontsize=8, ha="center", va="center",
                color="#444444" if external else "#111111")

    def arrow(x1, y1, x2, y2, label=None, two_way=False, ls="-"):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="<->" if two_way else "->",
                                    color="#333333", lw=1.3, linestyle=ls))
        if label:
            ax.text((x1 + x2) / 2, (y1 + y2) / 2, label, fontsize=7,
                    color="#555555", ha="center", va="center",
                    bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="none"))

    # Hospital zone (top)
    zone(4, 55, 92, 42, "HOSPITAL NETWORK")
    box(22, 78, 26, 14, ["Imaging Unit", "(RTOS firmware)"])
    box(68, 78, 34, 14, ["Acquisition Workstation", "(Windows 10)",
                         "Scanning SW · Admin console :8443"])
    box(68, 61, 30, 9, ["Local DICOM Server", "(stores scan images)"])
    box(24, 61, 28, 9, ["Hospital EMR", "(external, partially-trusted)"],
        external=True)
    arrow(35, 78, 51, 78, "USB / proprietary", two_way=True)
    arrow(68, 71, 68, 65.5, "Hospital LAN")
    arrow(53, 74, 34, 65.8, "HL7 FHIR (TB-6)", two_way=True, ls="--")

    # Cloud zone (bottom)
    zone(4, 6, 92, 40, "MEDISCANTECH CLOUD")
    box(24, 34, 30, 10, ["AI Inference Service", "(REST API)"])
    box(66, 34, 28, 10, ["Cloud Storage", "(anon. images)"])
    box(50, 20, 66, 8, ["Manufacturer Backend (training, monitoring, support)"])
    box(50, 10, 66, 7, ["Update Delivery Service"])
    arrow(51, 34, 39, 34, two_way=True)

    # Cross-boundary flows (kept in the gap between the two zones to avoid
    # crossing any component boxes)
    arrow(34, 55, 30, 39.5, "Internet\n(TLS 1.3, mutual TLS)")
    arrow(84, 16, 84, 54.5, "Updates\n(signed pkgs)", ls="--")

    fig.tight_layout(pad=0.5)
    out = Path(tempfile.gettempdir()) / "neuroscan_architecture.png"
    fig.savefig(out, dpi=200, bbox_inches="tight")
    plt.close(fig)
    return out


def add_product_section(doc):
    """Section 2: the NeuroScan 3000 scenario (from scenario/*.md)."""
    doc.add_heading("Product & Architecture — NeuroScan 3000", level=1)
    add_hint(doc, "This is a fictional device created for training purposes. Any resemblance "
                  "to real products is coincidental.")

    # --- Overview ---
    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "NeuroScan 3000 is a neurological imaging system manufactured by MediScanTech Inc. "
        "It combines MRI-guided imaging with AI-assisted diagnostics to help clinicians detect "
        "neurological conditions such as stroke, tumours, and MS lesions.")
    doc.add_paragraph(
        "The system is hybrid: it has hardware on-premise in the hospital, and a cloud platform "
        "operated by MediScanTech that provides AI analysis and software updates. Both parts are "
        "essential to its operation.")

    doc.add_heading("Intended use", level=2)
    add_bullets(doc, [
        "Clinical setting: radiology departments in hospitals",
        "Primary users: radiologists, biomedical engineers, IT administrators",
        "Regulatory classification: Class II medical device (EU MDR Class IIb, FDA 510(k))",
    ])

    # --- Components ---
    doc.add_heading("System components", level=2)
    doc.add_heading("On-premise (hospital)", level=3)
    add_info_table(doc,
                   ["Component", "Description"],
                   [
                       ["Imaging unit", "Physical MRI scanner with embedded firmware "
                        "(proprietary operating system)"],
                       ["Acquisition workstation", "PC connected to the imaging unit. Runs the "
                        "scanning software and a local admin console accessible on the hospital "
                        "network."],
                       ["Local DICOM server", "Stores scan images on-site. Connected to the "
                        "hospital network."],
                   ])
    doc.add_heading("Cloud (MediScanTech-operated)", level=3)
    add_info_table(doc,
                   ["Component", "Description"],
                   [
                       ["AI inference service", "Receives anonymised scan images, runs the AI "
                        "diagnostic model, and returns annotations"],
                       ["Cloud storage", "Stores anonymised images and AI outputs"],
                       ["Update delivery service", "Delivers firmware and software updates to "
                        "on-premise components"],
                       ["Manufacturer backend", "MediScanTech internal systems: model training, "
                        "monitoring, remote support"],
                   ])

    # --- Data flow ---
    doc.add_heading("How a scan works (data flow)", level=2)
    add_numbered(doc, [
        "Patient undergoes scan → imaging unit captures raw data",
        "Acquisition workstation processes the raw data → creates DICOM images",
        "Images stored on the local DICOM server",
        "Anonymised images sent to the cloud AI inference service (encrypted)",
        "Cloud returns AI diagnostic annotations → displayed to radiologist on the workstation",
        "MediScanTech pushes firmware and software updates to the workstation and imaging unit",
    ])

    # --- Users ---
    doc.add_heading("Users and roles", level=2)
    add_info_table(doc,
                   ["Role", "What they do", "Where they access from"],
                   [
                       ["Radiologist", "Views images and AI annotations, signs diagnostic reports",
                        "Hospital workstation"],
                       ["Biomedical engineer", "Maintains the device, accesses local admin console",
                        "Hospital (on-site)"],
                       ["Hospital IT admin", "Manages network and local server", "Hospital network"],
                       ["MediScanTech support", "Remote access for diagnostics and troubleshooting",
                        "Internet"],
                       ["MediScanTech engineer", "Manages cloud backend and update delivery",
                        "MediScanTech datacentre"],
                   ])

    # --- Security assumptions ---
    doc.add_heading("Key security assumptions (read carefully)", level=2)
    add_bullets(doc, [
        "The acquisition workstation is connected to the internet — it talks to the cloud AI "
        "service and receives updates",
        "The hospital network is shared with other hospital systems (not a dedicated "
        "medical-only network)",
        "The local admin console has no multi-factor authentication and is reachable from "
        "anywhere on the hospital network",
        "Remote support uses a shared credential — one login used by all MediScanTech support "
        "staff (known gap)",
        "Anonymisation of images before cloud upload is done in software and has not been "
        "independently audited",
        "The imaging unit firmware is patched approximately every 18 months due to regulatory "
        "constraints",
        "The workstation integrates with the hospital EMR, which MediScanTech does not own or "
        "secure — treat it as a partially-trusted external dependency and do not assume data "
        "received from it is safe",
    ])

    # --- Architecture diagram ---
    doc.add_heading("Architecture diagram", level=2)
    add_hint(doc, "Boxes = components, dashed frames = trust zones, arrows = data flows.")
    img = render_architecture_png()
    doc.add_picture(str(img), width=Inches(6.3))
    # Center the image.
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Trust boundaries ---
    doc.add_heading("Trust boundaries", level=2)
    add_info_table(doc,
                   ["Boundary", "Between", "Notes"],
                   [
                       ["TB-1", "Imaging unit ↔ Acquisition workstation",
                        "Physical/proprietary — assumed trusted within the device"],
                       ["TB-2", "Acquisition workstation ↔ Hospital LAN",
                        "Shared hospital network — may include non-medical systems"],
                       ["TB-3", "Hospital LAN ↔ Local DICOM server",
                        "Internal — should be on a dedicated VLAN but often is not"],
                       ["TB-4", "Hospital ↔ MediScanTech cloud",
                        "Internet — protected by mutual TLS; main external boundary"],
                       ["TB-5", "MediScanTech backend ↔ Update service",
                        "Internal cloud — must be authenticated and integrity-protected"],
                       ["TB-6", "Acquisition workstation ↔ Hospital EMR",
                        "External system, hospital-owned — separate security governance; "
                        "treat as partially-trusted"],
                   ])

    # --- Data flows ---
    doc.add_heading("Data flows", level=2)
    add_info_table(doc,
                   ["Data", "Sensitivity", "Path"],
                   [
                       ["Raw DICOM images (with patient data)", "High",
                        "Imaging unit → Workstation → Local DICOM server"],
                       ["Anonymised DICOM images", "Medium",
                        "Local DICOM server → Cloud AI service → Cloud storage"],
                       ["AI diagnostic annotations", "Medium",
                        "Cloud → Workstation → Local DICOM server"],
                       ["Device logs / telemetry", "Low–Medium",
                        "Workstation → MediScanTech backend"],
                       ["Firmware / software updates", "High (integrity)",
                        "MediScanTech → Workstation + Imaging unit"],
                       ["Admin credentials", "Critical",
                        "Local admin console, remote support"],
                       ["Diagnostic reports / patient & order data", "High",
                        "Workstation ↔ Hospital EMR (HL7 FHIR)"],
                   ])

    # --- Interfaces ---
    doc.add_heading("Key interfaces", level=2)
    add_info_table(doc,
                   ["Interface", "Protocol", "Authentication", "Notes"],
                   [
                       ["Imaging unit → Workstation", "Proprietary USB", "None (physical)",
                        "Assumed trusted"],
                       ["Workstation → Local DICOM server", "DICOM (TCP 104)", "None by default",
                        "No built-in auth in most deployments"],
                       ["Workstation → Cloud AI service", "HTTPS REST", "Mutual TLS + API key", ""],
                       ["Local admin console", "HTTPS (:8443)", "Username/password", "No MFA"],
                       ["Remote support", "VPN + SSH", "Shared credential", "Known security gap"],
                       ["Update delivery", "HTTPS", "Package signing (RSA-2048) + TLS", ""],
                       ["Workstation → Hospital EMR", "HL7 FHIR over HTTPS", "OAuth 2.0",
                        "External hospital-owned system; treat as partially-trusted"],
                   ])


def main():
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else (
        Path(__file__).resolve().parents[4] / "templates" / "threat-model-template.docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style_base(doc)

    # ---- Section 1: Product & architecture ----
    # (The workshop introduction lives in workshop/00-introduction.md and is not
    #  duplicated here — the facilitator shares it from the repo.)
    add_product_section(doc)
    page_break(doc)

    # ---- Section 2: Threat model worksheet ----
    doc.add_heading("Threat Model Worksheet — NeuroScan 3000", level=1)
    meta = doc.add_paragraph()
    meta.add_run("Team name: ").bold = True
    meta.add_run("____________________________     ")
    meta.add_run("Date: ").bold = True
    meta.add_run("____________________")
    doc.add_paragraph("Workshop: Medical Device Threat Modeling")
    add_hint(doc, "Fill this worksheet in as you work through Q1–Q4. Import it into "
                  "Google Docs (File → Open → Upload), complete it as a group, then "
                  "export back to .docx or PDF for AI evaluation.")

    # ---- Q1 ----
    doc.add_heading("Q1: What are we working on?", level=2)
    add_hint(doc, "Goal: build a shared map of the system before you look for threats. "
                  "Agree what's in scope, what's worth protecting, and who could cause harm.")

    doc.add_heading("1.1 System boundary", level=3)
    doc.add_paragraph("In scope — components and interfaces your team will analyze:")
    blank_lines(doc, 3)
    doc.add_paragraph("Out of scope — what you exclude, and why "
                      "(e.g. \"hospital Wi-Fi — managed by hospital IT, outside MediScanTech's control\"):")
    blank_lines(doc, 2)

    doc.add_heading("1.2 Assets — what are we protecting?", level=3)
    add_hint(doc, "Ask: what would hurt if it was stolen, changed, or made unavailable? "
                  "Property key: C = Confidentiality, I = Integrity, A = Availability. Aim for ≥5.")
    add_table(doc,
              ["Asset ID", "Asset", "Why it matters", "Most critical property (C/I/A)"],
              n_blank_rows=5,
              example_row=["A-01", "DICOM scan images with patient data",
                           "Contains PHI; misuse could harm patients or breach privacy", "Confidentiality"])

    doc.add_heading("1.3 Entry points — how can someone get in?", level=3)
    add_hint(doc, "Network ports, physical connections, web portals, update channels, remote access.")
    add_table(doc,
              ["Entry point", "How it works", "Authentication?", "Encrypted?", "Who can reach it?"],
              n_blank_rows=5,
              example_row=["Local admin console (:8443)", "Web UI on the workstation",
                           "Username/password, no MFA", "Yes (HTTPS)", "Anyone on hospital LAN"])

    doc.add_heading("1.4 Threat actors — who might attack this?", level=3)
    add_hint(doc, "Go beyond opportunistic hackers — insiders, nation-states, ransomware groups, vendor staff.")
    add_table(doc,
              ["Actor", "Motivation", "How they might get in"],
              n_blank_rows=5,
              example_row=["Ransomware group", "Financial — encrypt systems, demand payment",
                           "Phishing hospital staff; exploiting internet-facing services"])

    # ---- Q2 ----
    doc.add_heading("Q2: What can go wrong?", level=2)
    add_hint(doc, "Goal: think like an attacker. Write stories from the attacker's perspective, "
                  "then assess how serious each one is.")

    doc.add_heading("2.1 Attacker stories", level=3)
    p = doc.add_paragraph()
    p.add_run("Format: ").bold = True
    p.add_run("As a [bad actor], I want to [do something bad] via [method or entry point], "
              "so that [I achieve my goal].").italic = True
    add_hint(doc, "Example: As a rogue vendor technician, I want to access the system beyond my "
                  "support role via the shared remote support credential, so that I can exfiltrate "
                  "patient data or plant a backdoor. Aim for ≥8 stories across different parts of the system.")
    add_table(doc,
              ["Story ID", "Bad actor", "Attacker story", "Part of system affected", "Impact type (S/P/A)"],
              n_blank_rows=8,
              example_row=["S-01", "Ransomware group",
                           "As a ransomware group, I want to encrypt the acquisition workstation via a "
                           "phishing email, so that the hospital cannot perform scans and must pay a ransom.",
                           "Acquisition workstation", "A"])
    add_hint(doc, "Impact type: S = patient safety (physical harm) · P = privacy / PHI "
                  "(confidentiality) · A = availability (loss of scanning / care disruption). "
                  "Tag the dominant type — it tells you which harm to score in 2.2. A threat can carry more than one.")

    doc.add_heading("2.2 Risk assessment", level=3)
    p = doc.add_paragraph()
    p.add_run("Exploitability: ").bold = True
    p.add_run("1 = Low (rare access / high skill) · 2 = Moderate (remote, known weakness) · "
              "3 = High (trivial, public exploit)")
    p1b = doc.add_paragraph()
    p1b.add_run("Severity of harm (patient safety first — also privacy & availability): ").bold = True
    p1b.add_run("1 = Minor (no patient harm; negligible privacy/operational impact) · "
                "2 = Significant (care delayed/degraded, a privacy/PHI breach, or meaningful downtime) · "
                "3 = Serious (direct patient harm, a large-scale PHI breach, or extended loss of scanning)")
    p2 = doc.add_paragraph()
    p2.add_run("Risk = Exploitability × Severity  →  1–2: Low · 3–4: Medium · 6–9: High").bold = True
    note = doc.add_paragraph()
    nr = note.add_run("Patient safety rule: if a story could directly harm a patient, mark it High "
                      "regardless of score and add a note. Score a confidentiality-only threat on the "
                      "severity of its privacy harm — not \"no harm\".")
    nr.bold = True
    fda = doc.add_paragraph()
    fr = fda.add_run("FDA note: security risk is scored on exploitability (how feasible the attack "
                     "is), not probability — the FDA does not permit probabilistic scoring for "
                     "security risk. Exploitability × severity is the FDA's controlled/uncontrolled matrix.")
    fr.italic = True
    fr.font.color.rgb = NOTE
    fr.font.size = Pt(10)
    eng = doc.add_paragraph()
    er = eng.add_run("In a real engagement: this 3×3 is a teaching scale. Production threat models "
                     "usually layer STRIDE (systematic enumeration per component/data flow) and CVSS "
                     "(standardized exploitability scoring) on top — noting base CVSS doesn't capture "
                     "patient-safety severity, so it's adapted with a medical-device rubric, not used raw.")
    er.italic = True
    er.font.color.rgb = NOTE
    er.font.size = Pt(10)
    add_table(doc,
              ["Story ID", "Exploitability (1–3)", "Rationale", "Severity (1–3)", "Rationale",
               "Risk score", "Patient safety?", "Priority"],
              n_blank_rows=8)

    # ---- Q3 ----
    doc.add_heading("Q3: What are we going to do about it?", level=2)
    add_hint(doc, "Goal: for each high-priority story, define at least one concrete mitigation. "
                  "Be specific. Note trade-offs and residual risk.")
    p = doc.add_paragraph()
    p.add_run("Types: ").bold = True
    p.add_run("Preventive (stops it) · Detective (spots it) · Corrective (limits damage) · "
              "Compensating (alternative when the ideal fix isn't feasible)")
    add_hint(doc, "\"Add encryption\" is not specific enough. \"Enable TLS on the DICOM connection "
                  "between the workstation and the local DICOM server\" is.")
    add_table(doc,
              ["Mitigation ID", "Addresses story(ies)", "Type", "What exactly would be done",
               "Residual risk", "Regulatory note"],
              n_blank_rows=6,
              example_row=["M-01", "S-03, S-05", "Preventive",
                           "Replace shared remote support credential with per-engineer certificates; log all sessions",
                           "Insider threat reduced, not eliminated", "Config change only — no new submission"])

    # ---- Q4 ----
    doc.add_heading("Q4: Did we do a good job?", level=2)
    add_hint(doc, "Goal: step back and review your work critically before presenting.")

    doc.add_heading("Our top 3 threats", level=3)
    for i in (1, 2, 3):
        doc.add_paragraph(f"{i}.  S-____:  ______________________________________________________")
    doc.add_heading("Most surprising finding", level=3)
    blank_lines(doc, 2)
    doc.add_heading("Hardest trade-off", level=3)
    blank_lines(doc, 2)
    doc.add_heading("What we're still unsure about", level=3)
    blank_lines(doc, 2)

    doc.add_heading("Completeness checklist", level=3)
    checklist = [
        "At least 8 attacker stories, each covering a different part of the system",
        "Every story written in the full format (actor → action → method → goal)",
        "Every story has an exploitability, severity, and risk score",
        "Any story that could harm a patient is marked High priority with rationale",
        "Every High-priority story has at least one mitigation",
        "Every Medium-priority story has a mitigation or a documented acceptance rationale",
        "Every mitigation says specifically what would be done (not just the category)",
        "Residual risk is noted for each mitigation",
        "Assets and entry points carry stable IDs (supports FDA traceability to the SBOM)",
    ]
    for item in checklist:
        doc.add_paragraph(f"☐  {item}")

    doc.save(str(out))
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
