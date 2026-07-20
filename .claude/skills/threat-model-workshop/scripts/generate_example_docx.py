#!/usr/bin/env python3
"""
Render the filled reference example to a Google-Docs-friendly .docx.

Source of truth is templates/example-filled-template.md (the completed
"Facilitators — Reference Example" threat model that uses attacker stories). This
script renders that markdown to templates/example-filled-template.docx so
facilitators can hand out or upload the worked example in the same format teams
use, and so it stays in sync with the markdown rather than being maintained twice.

It reuses the styling helpers from generate_template_docx.py (same fonts, heading
styles, and shaded table headers) so the example looks like the blank worksheet.

Supported markdown (only what the example uses): H1/H2/H3 headings, blockquotes,
pipe tables, bullet lists (-), numbered lists (1.), horizontal rules (---), blank
lines, and inline **bold**.

Requirements: python-docx.

Usage:
    python generate_example_docx.py [output_path]
Default output: ../../../templates/example-filled-template.docx
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from generate_template_docx import (  # noqa: E402
    HEADER_FILL, NOTE, add_hint, add_info_table, set_cell_background, style_base,
)

SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parents[3]
DEFAULT_SRC = REPO_ROOT / "templates" / "example-filled-template.md"
DEFAULT_OUT = REPO_ROOT / "templates" / "example-filled-template.docx"

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_runs_with_bold(paragraph, text):
    """Add text to a paragraph, honouring **bold** spans."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def strip_inline(text):
    """Plain text with markdown emphasis markers removed (for table cells)."""
    return BOLD_RE.sub(r"\1", text).strip()


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line):
    return bool(re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", line)) and "-" in line


def add_bold_header_table(doc, headers, rows):
    """Content table with a shaded header row; cell **bold** is honoured."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(strip_inline(h))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_background(cell, HEADER_FILL)
    for i, row in enumerate(rows, start=1):
        for j in range(len(headers)):
            val = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = ""
            add_runs_with_bold(cell.paragraphs[0], val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def render_markdown(md_text, doc):
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(strip_inline(m.group(2)), level=level)
            i += 1
            continue

        # Table (header row followed by a separator row)
        if stripped.startswith("|") and i + 1 < n and is_table_separator(lines[i + 1]):
            headers = parse_table_row(lines[i])
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_bold_header_table(doc, headers, rows)
            continue

        # Blockquote (may span multiple lines); render as an indented hint/note
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            text = " ".join(q.strip() for q in quote_lines if q.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs_with_bold(p, text)
            for r in p.runs:
                r.font.color.rgb = NOTE
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", stripped):
            while i < n and re.match(r"^[-*]\s+", lines[i].strip()):
                item = re.sub(r"^[-*]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_runs_with_bold(p, item)
                i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_runs_with_bold(p, item)
                i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1


def main():
    src = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_SRC
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUT
    out.parent.mkdir(parents=True, exist_ok=True)

    md_text = src.read_text(encoding="utf-8")
    doc = Document()
    style_base(doc)
    render_markdown(md_text, doc)
    doc.save(str(out))
    print(f"Wrote {out}  (from {src.name})")


if __name__ == "__main__":
    main()
