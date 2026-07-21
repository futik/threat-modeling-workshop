#!/usr/bin/env python3
"""
Generate sample participant submissions (filled worksheets) as .docx files.

This produces the kind of files teams hand back after the workshop: the Q1-Q4
worksheet, filled in, exported as .docx. It writes one .docx per team into the
examples/submissions/ folder (the committed worked example), so the
batch-evaluation workflow (see SKILL.md, Mode B) has realistic inputs to iterate
over. Real participant submissions go in submissions/ instead.

The teams deliberately vary in quality so the evaluation and cross-team
comparison have something to differentiate:
  - Team Aegis      — strong (broad scope, exploitability-based scoring, specific
                      mitigations, FDA/lifecycle awareness)
  - Team Meridian   — medium (decent coverage, some vague mitigations, partial
                      patient-safety overrides, thin regulatory awareness)
  - Team Northwind  — weak (narrow scope, generic threats, probabilistic scoring,
                      "add encryption"-style mitigations, no regulatory context)

The .docx layout matches the fillable worksheet in generate_template_docx.py so
extract_submission.py reads it back the same way it reads real submissions.

Usage:
    python generate_sample_submissions.py [output_dir]
Default output dir: ../../../examples/submissions
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

# Reuse the styling helpers from the template generator so submissions look like
# genuine exports of the worksheet.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from generate_template_docx import set_cell_background, style_base, HEADER_FILL  # noqa: E402

NOTE = RGBColor(0x55, 0x55, 0x55)


def filled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_background(cell, HEADER_FILL)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def build_submission(team, out_path):
    doc = Document()
    style_base(doc)

    doc.add_heading("Threat Model Worksheet — NeuroScan 3000", level=1)
    meta = doc.add_paragraph()
    meta.add_run("Team name: ").bold = True
    meta.add_run(team["name"] + "     ")
    meta.add_run("Date: ").bold = True
    meta.add_run(team["date"])
    doc.add_paragraph("Workshop: Medical Device Threat Modeling")

    # ---- Q1 ----
    doc.add_heading("Q1: What are we working on?", level=2)

    doc.add_heading("1.1 System boundary", level=3)
    doc.add_paragraph("In scope — components and interfaces your team will analyze:")
    for item in team["in_scope"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("Out of scope — what you exclude, and why:")
    for item in team["out_scope"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.2 Assets — what are we protecting?", level=3)
    filled_table(doc,
                 ["Asset ID", "Asset", "Why it matters", "Most critical property (C/I/A)"],
                 team["assets"])

    doc.add_heading("1.3 Entry points — how can someone get in?", level=3)
    filled_table(doc,
                 ["Entry point", "How it works", "Authentication?", "Encrypted?",
                  "Who can reach it?"],
                 team["entry_points"])

    doc.add_heading("1.4 Threat actors — who might attack this?", level=3)
    filled_table(doc,
                 ["Actor", "Motivation", "How they might get in"],
                 team["actors"])

    # ---- Q2 ----
    doc.add_heading("Q2: What can go wrong?", level=2)

    doc.add_heading("2.1 Attacker stories", level=3)
    filled_table(doc,
                 ["Story ID", "Bad actor", "Attacker story", "Part of system affected",
                  "Impact type (S/P/A)"],
                 team["stories"])

    doc.add_heading("2.2 Risk assessment", level=3)
    filled_table(doc,
                 ["Story ID", "Exploitability (1–3)", "Rationale", "Severity (1–3)",
                  "Rationale", "Risk score", "Patient safety?", "Priority"],
                 team["risk"])

    # ---- Q3 ----
    doc.add_heading("Q3: What are we going to do about it?", level=2)
    filled_table(doc,
                 ["Mitigation ID", "Addresses story(ies)", "Type",
                  "What exactly would be done", "Residual risk", "Regulatory note"],
                 team["mitigations"])

    # ---- Q4 ----
    doc.add_heading("Q4: Did we do a good job?", level=2)
    doc.add_heading("Our top 3 threats", level=3)
    for line in team["top3"]:
        doc.add_paragraph(line)
    doc.add_heading("Most surprising finding", level=3)
    doc.add_paragraph(team["surprising"])
    doc.add_heading("Hardest trade-off", level=3)
    doc.add_paragraph(team["tradeoff"])
    doc.add_heading("What we're still unsure about", level=3)
    for line in team["unsure"]:
        doc.add_paragraph(line, style="List Bullet")

    doc.save(str(out_path))
    return out_path


# --------------------------------------------------------------------------- #
# Team data
# --------------------------------------------------------------------------- #

TEAM_AEGIS = {
    "name": "Team Aegis",
    "date": "2026-05-14",
    "slug": "team-aegis",
    "in_scope": [
        "Imaging unit (MRI hardware + embedded RTOS firmware)",
        "Acquisition workstation (Windows 10, scanning software, admin console :8443)",
        "Local DICOM server",
        "Cloud AI inference service and cloud storage (anonymised images)",
        "Update delivery service and manufacturer backend",
        "All data flows and interfaces between the above (TB-1 through TB-5)",
    ],
    "out_scope": [
        "Hospital EMR/HIS — separate procurement and governance; modelled as a "
        "partially-trusted external system (boundary TB-6). We do not own it, but we "
        "validate all data crossing the EMR interface as untrusted input",
        "Hospital network hardware (switches, firewalls) — outside MediScanTech control; "
        "treated as the threat boundary rather than a controlled component",
    ],
    "assets": [
        ["A-01", "Patient DICOM images with PHI", "Contains PHI; disclosure harms patients and "
         "breaches privacy law", "Confidentiality + Integrity"],
        ["A-02", "AI diagnostic annotations", "Clinical decision support; tampering can cause "
         "misdiagnosis and patient harm", "Integrity"],
        ["A-03", "Imaging unit firmware", "Compromise gives persistent control and can alter "
         "scan behaviour", "Integrity"],
        ["A-04", "Firmware/software update packages", "A malicious update compromises every "
         "fielded device", "Integrity"],
        ["A-05", "Admin + remote-support credentials", "Full control of the device if stolen",
         "Confidentiality"],
        ["A-06", "Scan availability (ability to image)", "Downtime delays diagnosis in "
         "time-critical cases (e.g. stroke)", "Availability"],
        ["A-07", "Anonymisation logic", "If it under-scrubs, PHI leaks to the cloud",
         "Confidentiality + Integrity"],
    ],
    "entry_points": [
        ["Local admin console (:8443)", "Web UI on the workstation", "Username/password, no MFA",
         "Yes (HTTPS)", "Anyone on the hospital LAN"],
        ["Remote support (VPN+SSH)", "Vendor remote access", "Shared credential, no MFA",
         "Yes (VPN)", "MediScanTech support staff (internet)"],
        ["Workstation ↔ Local DICOM server", "DICOM TCP/104", "None by default", "None by default",
         "Anyone on the hospital LAN"],
        ["Workstation ↔ Cloud AI service", "HTTPS REST", "Mutual TLS + API key", "TLS 1.3",
         "Internet-facing"],
        ["Update delivery", "Signed package pull over HTTPS", "Package signing (RSA-2048)",
         "TLS 1.3", "Internet-facing"],
        ["Workstation ↔ Hospital EMR (TB-6)", "HL7 FHIR over HTTPS", "OAuth 2.0", "TLS 1.3",
         "Partially-trusted external system"],
        ["Imaging unit ↔ Workstation", "Proprietary USB", "None (physical)", "None",
         "Physical access on-site"],
    ],
    "actors": [
        ["Ransomware group", "Financial — encrypt systems, demand payment",
         "Phishing hospital staff; exploiting internet-facing services"],
        ["Malicious hospital insider", "Sabotage, data theft, bribery",
         "Legitimate LAN access and workstation credentials"],
        ["Nation-state actor", "Healthcare disruption, espionage",
         "Supply-chain compromise, spear-phishing, zero-days"],
        ["Rogue MediScanTech engineer", "Financial, espionage",
         "Privileged access to cloud backend and update delivery"],
        ["Opportunistic external attacker", "Data resale, notoriety",
         "Man-in-the-middle on the hospital LAN; brute force of exposed services"],
    ],
    "stories": [
        ["S-01", "Ransomware group", "As a ransomware group, I want to encrypt the acquisition "
         "workstation via a phishing email to hospital staff, so that scanning halts and the "
         "hospital pays a ransom.", "Acquisition workstation", "A"],
        ["S-02", "Opportunistic attacker", "As an attacker with a LAN foothold, I want to modify "
         "unencrypted DICOM traffic between the workstation and the local DICOM server, so that "
         "the radiologist diagnoses from corrupted images.", "Workstation ↔ DICOM (TB-3)", "S"],
        ["S-03", "External attacker", "As an attacker, I want to brute-force the local admin "
         "console (no MFA, no rate limiting) via the hospital LAN, so that I gain workstation "
         "administrator access.", "Admin console :8443", "S"],
        ["S-04", "Rogue vendor technician", "As a support technician, I want to use the shared "
         "remote-support credential beyond my role, so that I can exfiltrate PHI or plant a "
         "backdoor undetected.", "Remote support", "P"],
        ["S-05", "Nation-state actor", "As a nation-state, I want to push a malicious firmware "
         "update via the update delivery service, so that I persistently control fielded imaging "
         "units and can manipulate scans.", "Update delivery → firmware", "S"],
        ["S-06", "Rogue MediScanTech engineer", "As a cloud engineer, I want to alter the AI "
         "model weights, so that the service systematically misclassifies specific lesion types "
         "across all hospitals.", "Cloud AI model", "S"],
        ["S-07", "Malicious insider", "As an insider with cloud access, I want to re-identify "
         "anonymised images using external demographic data, so that I can breach patient "
         "privacy at scale.", "Cloud storage", "P"],
        ["S-08", "External attacker", "As an attacker, I want to flood the cloud AI inference "
         "service, so that annotations are unavailable during active scans and the workflow is "
         "disrupted.", "Cloud AI service", "A"],
        ["S-09", "Opportunistic attacker", "As an attacker, I want to exploit weak anonymisation "
         "(burned-in PHI in pixel data), so that PHI is transmitted to the cloud in the clear.",
         "Anonymisation boundary", "P"],
        ["S-10", "Malicious insider", "As an admin-level attacker, I want to clear the audit "
         "logs on the workstation and DICOM server, so that my actions cannot be attributed.",
         "Workstation + DICOM logs", "P"],
        ["S-11", "External attacker (via hospital EMR)", "As an attacker who has compromised the "
         "partially-trusted hospital EMR, I want to send malformed HL7 FHIR messages across the "
         "EMR interface to the workstation, so that I exploit a parser flaw to run code or feed "
         "corrupted patient/order data into a scan.", "Workstation ↔ EMR (TB-6)", "S"],
    ],
    "risk": [
        ["S-01", "3", "Phishing is trivial; workstation is internet-connected", "3",
         "Extended loss of scanning; delays time-critical diagnosis", "9", "Yes", "High"],
        ["S-02", "2", "LAN foothold achievable; DICOM often unencrypted", "3",
         "Corrupted images drive misdiagnosis — direct patient harm", "6", "Yes", "High"],
        ["S-03", "3", "No MFA, no rate limiting, whole-LAN reachable", "3",
         "Full workstation control; pivot to connected systems", "9", "Yes", "High"],
        ["S-04", "2", "Shared credential known to many; no MFA", "3",
         "Undetectable full access; PHI loss and backdoor", "6", "Yes", "High"],
        ["S-05", "1", "Requires compromising the signed update pipeline", "3",
         "Persistent fleet compromise; scan manipulation", "3", "Yes",
         "High (safety override)"],
        ["S-06", "1", "Needs privileged model-deploy access; hard to conceal", "3",
         "Systematic misdiagnosis across hospitals", "3", "Yes", "High (safety override)"],
        ["S-07", "1", "Sophisticated insider + external data needed", "2",
         "Large-scale PHI breach; regulatory exposure", "2", "No", "Medium"],
        ["S-08", "2", "Public endpoint; may lack rate limiting", "2",
         "Workflow disruption; radiologist can proceed without AI", "4", "No", "Medium"],
        ["S-09", "2", "Anonymisation bugs are a known DICOM issue class", "2",
         "PHI breach; GDPR exposure", "4", "No", "Medium"],
        ["S-10", "2", "Needs prior admin access; log clearing then trivial", "1",
         "Forensics impaired; no direct harm alone", "2", "No", "Low"],
        ["S-11", "2", "EMR is partially-trusted and outside our control; HL7/FHIR parser flaws "
         "are a known class", "3", "Code execution on the workstation or corrupted order data "
         "used during a scan", "6", "Yes", "High"],
    ],
    "mitigations": [
        ["M-01", "S-03, S-04", "Preventive", "Enforce TOTP MFA on the admin console; replace the "
         "shared remote-support credential with per-engineer PKI certificates; rate-limit logins "
         "(5/min, 30-min lockout) and log all sessions", "Credential theft still possible on a "
         "compromised host", "Software/config change — likely no new 510(k)"],
        ["M-02", "S-02", "Preventive", "Enable DICOM-over-TLS (port 2762) on all local DICOM "
         "connections; allowlist workstation IP/MAC", "A compromised workstation can still send "
         "malicious data once allowlisted", "Config change — document in security file"],
        ["M-03", "S-05, S-06", "Preventive + Detective", "Vendor-sign AI inference responses "
         "(JWS) and verify on the workstation before display; log all model-version transitions; "
         "add post-update firmware integrity attestation", "Signing-key compromise still allows "
         "tampering", "May be a design change — assess FDA impact"],
        ["M-04", "S-01", "Corrective", "Immutable off-site backups of the workstation image + "
         "tested restore runbook; EDR with isolation", "Ransom pressure remains during restore "
         "window", "Software change; no new submission"],
        ["M-05", "S-09", "Detective", "Integrate a DICOM anonymisation audit tool that scans for "
         "burned-in PHI before upload and blocks/loggs failures", "Zero-day anonymisation gaps "
         "missed until the tool updates", "Add to design verification"],
        ["M-06", "S-08", "Compensating", "Graceful degradation: if AI is unavailable >60s, warn "
         "the radiologist and allow manual-only review; alert on-call", "No AI assistance during "
         "degradation", "Software change"],
        ["M-07", "S-11", "Preventive", "Treat the EMR interface (TB-6) as untrusted input: strict "
         "HL7 FHIR schema validation and sanitisation, run the FHIR parser in a sandboxed "
         "least-privilege process, and mutually authenticate the connection (OAuth 2.0 + pinned "
         "TLS)", "Zero-day parser flaws remain until patched; a fully compromised EMR can still "
         "send well-formed but false data", "Software change — add interface fuzzing to design "
         "verification"],
    ],
    "top3": [
        "1.  S-03:  No MFA on the admin console with whole-LAN reach — trivial full compromise",
        "2.  S-01:  Ransomware on the internet-connected workstation halts scanning",
        "3.  S-02:  Unencrypted DICOM lets an attacker corrupt images and cause misdiagnosis",
    ],
    "surprising": "S-06 — model-weight manipulation. The AI is a black box the hospital doesn't "
                  "control, so a systematic bias could harm thousands of patients before anyone "
                  "notices. It has no clean precedent in traditional device threat modelling.",
    "tradeoff": "Properly fixing S-05 needs secure boot with a hardware root of trust, which "
                "forces a hardware revision and a new FDA 510(k). We chose signed updates + "
                "post-update integrity attestation as an interim control while planning the "
                "hardware change, accepting the residual risk explicitly.",
    "unsure": [
        "MediScanTech's rotation SLA for the shared remote-support credential (S-04)",
        "Whether the anonymisation software (S-09) has ever been independently audited",
        "Whether the cloud AI service has its own threat model shared with hospitals",
    ],
}

TEAM_MERIDIAN = {
    "name": "Team Meridian",
    "date": "2026-05-14",
    "slug": "team-meridian",
    "in_scope": [
        "Acquisition workstation and admin console",
        "Local DICOM server",
        "Cloud AI inference service",
        "Update delivery",
    ],
    "out_scope": [
        "Hospital Wi-Fi — managed by hospital IT",
        "Imaging unit firmware — we assumed it was trusted (did not analyse in depth)",
    ],
    "assets": [
        ["A-01", "Patient images (PHI)", "Privacy and patient trust", "Confidentiality"],
        ["A-02", "AI results", "Used for diagnosis", "Integrity"],
        ["A-03", "Admin credentials", "Control of the workstation", "Confidentiality"],
        ["A-04", "Scanning availability", "Needed for patient care", "Availability"],
        ["A-05", "Update packages", "Could be abused to push bad code", "Integrity"],
    ],
    "entry_points": [
        ["Admin console :8443", "Web UI", "Password only", "HTTPS", "Hospital LAN"],
        ["Remote support", "Vendor login", "Shared credential", "VPN", "Internet"],
        ["Cloud AI API", "REST", "TLS + key", "Yes", "Internet"],
        ["DICOM server", "DICOM", "None", "No", "Hospital LAN"],
    ],
    "actors": [
        ["Hackers", "Money / ransomware", "Phishing, internet-facing services"],
        ["Insider", "Revenge or money", "Existing access"],
        ["Vendor staff", "Misuse of access", "Remote support credential"],
    ],
    "stories": [
        ["S-01", "Hackers", "As a hacker, I want to phish staff and deploy ransomware on the "
         "workstation, so that the hospital can't scan.", "Workstation", "A"],
        ["S-02", "Hackers", "As an attacker on the LAN, I want to read DICOM traffic, so that I "
         "can steal patient images.", "DICOM server", "P"],
        ["S-03", "Hackers", "As an attacker, I want to log into the admin console because there "
         "is no MFA, so that I control the workstation.", "Admin console", "P"],
        ["S-04", "Vendor staff", "As a support engineer, I want to use the shared credential to "
         "access more than I should, so that I can view patient data.", "Remote support", "P"],
        ["S-05", "Insider", "As an insider, I want to change AI results, so that a patient is "
         "misdiagnosed.", "Cloud AI service", "S"],
        ["S-06", "Hackers", "As an attacker, I want to push a fake update, so that I can install "
         "malware on the device.", "Update delivery", "S"],
        ["S-07", "Hackers", "As an attacker, I want to DoS the cloud AI, so that annotations are "
         "unavailable.", "Cloud AI service", "A"],
        ["S-08", "Insider", "As an insider, I want to copy anonymised images and re-identify "
         "them, so that I breach privacy.", "Cloud storage", "P"],
    ],
    "risk": [
        ["S-01", "3", "Phishing is easy", "3", "No scanning", "9", "Yes", "High"],
        ["S-02", "2", "LAN access needed", "2", "Privacy breach", "4", "No", "Medium"],
        ["S-03", "3", "No MFA", "2", "Attacker gets in", "6", "No", "High"],
        ["S-04", "2", "Shared credential", "2", "Data access", "4", "No", "Medium"],
        ["S-05", "2", "Needs access", "3", "Wrong diagnosis", "6", "Yes", "High"],
        ["S-06", "2", "Updates are signed but process is weak", "3", "Malware on device", "6",
         "Yes", "High"],
        ["S-07", "2", "Public API", "1", "AI down but manual works", "2", "No", "Low"],
        ["S-08", "1", "Hard to re-identify", "2", "Privacy breach", "2", "No", "Low"],
    ],
    "mitigations": [
        ["M-01", "S-03", "Preventive", "Add MFA to the admin console", "Some risk remains", ""],
        ["M-02", "S-02", "Preventive", "Encrypt DICOM traffic", "", ""],
        ["M-03", "S-04", "Preventive", "Give each engineer their own login instead of a shared "
         "one", "Insider risk remains", "Process change"],
        ["M-04", "S-01", "Corrective", "Keep backups so we can restore after ransomware",
         "Downtime during restore", ""],
        ["M-05", "S-05, S-06", "Preventive", "Sign AI results and updates and check the "
         "signatures", "Key compromise", ""],
    ],
    "top3": [
        "1.  S-01:  Ransomware stops scanning",
        "2.  S-03:  No MFA on the admin console",
        "3.  S-05:  AI result tampering leads to misdiagnosis",
    ],
    "surprising": "That the AI results could be tampered with — we first thought only data theft "
                  "mattered, but changing a diagnosis is worse.",
    "tradeoff": "MFA is good but engineers said it slows down emergency access, so we weren't "
                "sure how strict to make it.",
    "unsure": [
        "How firmware updates are actually verified on the imaging unit",
        "Whether encrypting DICOM breaks anything in the hospital",
    ],
}

TEAM_NORTHWIND = {
    "name": "Team Northwind",
    "date": "2026-05-14",
    "slug": "team-northwind",
    "in_scope": [
        "The workstation",
        "The cloud",
    ],
    "out_scope": [
        "Everything else",
    ],
    "assets": [
        ["A-01", "Patient data", "It's private", "Confidentiality"],
        ["A-02", "The system", "We need it to work", "Availability"],
        ["A-03", "Passwords", "Security", "Confidentiality"],
    ],
    "entry_points": [
        ["Website / console", "Login page", "Password", "Yes", "Users"],
        ["Internet", "Cloud connection", "Yes", "Yes", "Anyone"],
    ],
    "actors": [
        ["Hackers", "To steal data", "The internet"],
        ["Viruses", "Damage", "Email"],
    ],
    "stories": [
        ["S-01", "Hackers", "A hacker could steal patient data.", "Cloud", "P"],
        ["S-02", "Hackers", "A hacker could get the password and log in.", "Console", "P"],
        ["S-03", "Virus", "A virus could infect the workstation and break it.", "Workstation",
         "A"],
        ["S-04", "Hackers", "A hacker could take down the website.", "Cloud", "A"],
        ["S-05", "Hackers", "Someone could listen to the network traffic.", "Network", "P"],
    ],
    "risk": [
        ["S-01", "High", "Hackers are common", "High", "Data is sensitive", "High", "", "High"],
        ["S-02", "Medium", "Passwords can be guessed", "Medium", "They get in", "Medium", "",
         "Medium"],
        ["S-03", "High", "Viruses are everywhere", "High", "System breaks", "High", "", "High"],
        ["S-04", "Low", "DDoS is hard", "Medium", "Website down", "Low", "", "Low"],
        ["S-05", "Medium", "Sniffing is possible", "Low", "Maybe see data", "Low", "", "Low"],
    ],
    "mitigations": [
        ["M-01", "S-01", "Preventive", "Add encryption", "", ""],
        ["M-02", "S-02", "Preventive", "Use strong passwords", "", ""],
        ["M-03", "S-03", "Preventive", "Install antivirus", "", ""],
        ["M-04", "S-04", "Preventive", "Use a firewall", "", ""],
    ],
    "top3": [
        "1.  S-01:  Data theft",
        "2.  S-03:  Virus",
        "3.  S-02:  Password stolen",
    ],
    "surprising": "That there are so many ways to attack a medical device.",
    "tradeoff": "Encryption can slow things down.",
    "unsure": [
        "How the cloud works exactly",
    ],
}

TEAMS = [TEAM_AEGIS, TEAM_MERIDIAN, TEAM_NORTHWIND]


def main():
    out_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else (
        Path(__file__).resolve().parents[4] / "examples" / "submissions"
    )
    out_dir.mkdir(parents=True, exist_ok=True)
    for team in TEAMS:
        out_path = out_dir / f"threat-model-{team['slug']}.docx"
        build_submission(team, out_path)
        print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
